from datetime import datetime
from io import BytesIO

from docx import Document

SECTION_ORDER = [
    ("CONTEXT", "3. Контекст"), ("PROBLEM", "4. Проблема"), ("GOAL", "5. Цель"), ("BUSINESS_EFFECT", "6. Бизнес-эффект"),
    ("AS_IS", "7. AS IS"), ("TO_BE", "8. TO BE"), ("USE_CASES", "9. Use Cases"), ("FUNCTIONAL_REQUIREMENTS", "10. Функциональные требования"),
    ("NON_FUNCTIONAL_REQUIREMENTS", "11. Нефункциональные требования"), ("RISKS", "12. Риски"), ("FINAL_BT", "13. Финальный БТ"),
]

def _goal_section(doc: Document, sc: dict):
    doc.add_paragraph(f"Цель проекта: {sc.get('title') or 'Не заполнено'}")
    doc.add_paragraph("KPI / метрики успеха:")
    metrics = sc.get('successMetrics') or []
    if metrics:
        for m in metrics:
            doc.add_paragraph(f"- {m.get('metric','Не заполнено')} | текущее: {m.get('currentValue','Не заполнено')} | целевое: {m.get('targetValue','Не заполнено')}")
    else:
        doc.add_paragraph('Не заполнено')
    doc.add_paragraph('Что не входит в scope: ' + (', '.join(sc.get('nonGoals') or []) or 'Не заполнено'))
    doc.add_paragraph('Ограничения: ' + (', '.join(sc.get('constraints') or []) or 'Не заполнено'))
    doc.add_paragraph('Предпосылки: ' + (', '.join(sc.get('assumptions') or []) or 'Не заполнено'))
    doc.add_paragraph('Связанные проблемы: ' + (', '.join(sc.get('linkedProblems') or []) or 'Не заполнено'))

def build_docx(project, artifacts):
    doc = Document()
    doc.add_heading(project.project_name, 0)
    doc.add_paragraph(f"Статус: {project.status}")
    doc.add_paragraph(f"Дата: {datetime.utcnow().strftime('%Y-%m-%d')}")
    doc.add_paragraph("Версия: рабочая")
    doc.add_paragraph("Автор: Александр")
    doc.add_page_break()
    doc.add_heading("2. Содержание", level=1)
    for _, title in SECTION_ORDER:
        doc.add_paragraph(title)
    art = {a.artifact_type.value: a for a in artifacts}
    for key, title in SECTION_ORDER:
        doc.add_page_break(); doc.add_heading(title, level=1)
        a = art.get(key)
        if key == 'GOAL' and a and a.structured_content:
            _goal_section(doc, a.structured_content)
        else:
            doc.add_paragraph((a.content if a else "") or "Раздел пока не заполнен")
    bio = BytesIO(); doc.save(bio); bio.seek(0)
    return bio
