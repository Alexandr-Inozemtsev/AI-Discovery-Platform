from __future__ import annotations

import csv
from io import BytesIO, StringIO
from typing import Iterable

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {"txt", "md", "csv", "docx", "pdf", "xlsx"}
MAX_EXTRACTED_TEXT_CHARS = 60000
DEFAULT_CHUNK_SIZE = 3000
DEFAULT_MAX_CHUNKS = 12
MAX_XLSX_ROWS = 1000


def _decode_text(content: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return ""


def _normalize_text(text: str) -> str:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in normalized.split("\n")]
    return "\n".join(lines).strip()


def _trim_extracted_text(text: str) -> str:
    text = _normalize_text(text)
    if len(text) <= MAX_EXTRACTED_TEXT_CHARS:
        return text
    return text[:MAX_EXTRACTED_TEXT_CHARS].rstrip() + "\n...[truncated]"


def split_text_chunks(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    max_chunks: int = DEFAULT_MAX_CHUNKS,
) -> list[dict]:
    text = _normalize_text(text)
    if not text:
        return []
    chunks: list[dict] = []
    start = 0
    total = len(text)
    while start < total and len(chunks) < max_chunks:
        hard_end = min(start + chunk_size, total)
        end = hard_end
        if hard_end < total:
            boundary = max(text.rfind("\n", start, hard_end), text.rfind(" ", start, hard_end))
            if boundary > start + int(chunk_size * 0.6):
                end = boundary
        chunk_text = text[start:end].strip()
        if chunk_text:
            chunks.append(
                {
                    "order": len(chunks),
                    "text": chunk_text,
                    "metadata": {
                        "start": start,
                        "end": end,
                        "chars": len(chunk_text),
                        "truncated": end < total and len(chunks) + 1 >= max_chunks,
                    },
                }
            )
        start = max(end, start + 1)
        while start < total and text[start].isspace():
            start += 1
    return chunks


def _result(text: str, status: str | None = None, error: str | None = None) -> dict:
    text = _trim_extracted_text(text)
    final_status = status or ("completed" if text else "empty")
    if final_status != "completed":
        text = ""
    return {
        "text": text,
        "status": final_status,
        "error": error,
        "chunks": split_text_chunks(text) if final_status == "completed" else [],
    }


def _extract_docx(content: bytes) -> str:
    doc = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        txt = _normalize_text(paragraph.text)
        if txt:
            parts.append(txt)
    for table in doc.tables:
        for row in table.rows:
            cells = [_normalize_text(cell.text) for cell in row.cells]
            row_text = "\t".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages: list[str] = []
    for page_index, page in enumerate(reader.pages):
        page_text = _normalize_text(page.extract_text() or "")
        if page_text:
            pages.append(f"[Page {page_index + 1}]\n{page_text}")
    return "\n\n".join(pages)


def _iter_csv_rows(text: str) -> Iterable[str]:
    sample = text[:2048]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(StringIO(text), dialect)
    for row in reader:
        clean = [_normalize_text(cell) for cell in row]
        yield "\t".join(cell for cell in clean if cell)


def _extract_csv(content: bytes) -> str:
    decoded = _decode_text(content)
    return "\n".join(row for row in _iter_csv_rows(decoded) if row)


def _extract_xlsx(content: bytes) -> str:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    rows_seen = 0
    for sheet in workbook.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            if rows_seen >= MAX_XLSX_ROWS:
                parts.append("...[truncated rows]")
                return "\n".join(parts)
            values = [_normalize_text(str(value)) for value in row if value is not None and _normalize_text(str(value))]
            if values:
                parts.append("\t".join(values))
                rows_seen += 1
    return "\n".join(parts)


def _file_ext(filename: str) -> str:
    return (filename or "").lower().rsplit(".", 1)[-1] if "." in (filename or "") else ""


def extract_text_from_upload(filename: str, content: bytes, mime_type: str | None = None) -> dict:
    ext = _file_ext(filename)
    mime_type = (mime_type or "").lower()
    try:
        if ext in {"txt", "md"} or mime_type.startswith("text/plain"):
            return _result(_decode_text(content))
        if ext == "csv" or mime_type == "text/csv":
            return _result(_extract_csv(content))
        if ext == "docx":
            return _result(_extract_docx(content))
        if ext == "pdf":
            return _result(_extract_pdf(content))
        if ext == "xlsx":
            return _result(_extract_xlsx(content))
        if ext == "xls":
            return _result("", "unsupported", "Формат xls не поддержан. Сохраните файл как xlsx или csv.")
        return _result("", "unsupported", f"Неподдерживаемый формат файла: {ext or 'unknown'}")
    except Exception as exc:
        if ext not in SUPPORTED_EXTENSIONS and ext != "xls":
            return _result("", "unsupported", f"Неподдерживаемый формат файла: {ext or 'unknown'}")
        return _result("", "failed", f"Ошибка извлечения текста: {exc}")
